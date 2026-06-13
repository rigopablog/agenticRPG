import io
import os
import sys
import json
import asyncio
import requests
from pathlib import Path

from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from groq import Groq

sys.path.insert(0, str(Path(__file__).parent))
sys.path.insert(0, str(Path(__file__).parent.parent / "python"))
load_dotenv(Path(__file__).parent.parent / ".env")

from agent_defs import AGENTS, AGENTS_BY_ID

# Document libraries (imported at module level — listed in requirements.txt)
import pypdf
from docx import Document as DocxDocument
from fpdf import FPDF

app = FastAPI(title="Agents Hub")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://rpgdevelopment.com",
        "https://www.rpgdevelopment.com",
        "https://rpgdevelopment.net",
        "https://www.rpgdevelopment.net",
        "https://web-production-07e84.up.railway.app",
        "http://localhost:3000",
        "http://localhost:5173",
        "http://localhost:5500",
        "http://127.0.0.1:5500",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory=Path(__file__).parent / "static"), name="static")

groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
MODEL = "llama-3.3-70b-versatile"


# ── Request models ─────────────────────────────────────────────────────────────

class RunRequest(BaseModel):
    agent_id: str
    task: str
    context: str = ""
    document_context: str = ""  # text extracted from an uploaded file


class OrchestrationStep(BaseModel):
    agent_id: str
    task: str


class OrchestrationRequest(BaseModel):
    steps: list[OrchestrationStep]
    mode: str = "sequential"  # sequential | parallel


class DocRequest(BaseModel):
    content: str
    format: str = "pdf"   # "pdf" | "docx"
    filename: str = "reporte"


# ── Document helpers ──────────────────────────────────────────────────────────

def extract_pdf_text(content: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def extract_docx_text(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def generate_pdf_bytes(content: str, title: str = "Reporte") -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    for line in content.split("\n"):
        if not line.strip():
            pdf.ln(4)
        else:
            try:
                pdf.multi_cell(0, 6, line)
            except Exception:
                # Fallback: encode to latin-1 replacing unknown chars
                safe = line.encode("latin-1", errors="replace").decode("latin-1")
                pdf.multi_cell(0, 6, safe)
    return bytes(pdf.output())


def generate_docx_bytes(content: str, title: str = "Reporte") -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, 0)
    for line in content.split("\n"):
        doc.add_paragraph(line if line.strip() else "")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def root():
    html = (Path(__file__).parent / "static" / "index.html").read_text(encoding="utf-8")
    return HTMLResponse(html)


@app.get("/api/agents")
def get_agents():
    return AGENTS


@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    """Extract text from a PDF or DOCX file and return it as JSON."""
    content = await file.read()
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        try:
            text = extract_pdf_text(content)
        except Exception as e:
            raise HTTPException(422, f"No se pudo leer el PDF: {e}")
    elif ext in ("docx",):
        try:
            text = extract_docx_text(content)
        except Exception as e:
            raise HTTPException(422, f"No se pudo leer el DOCX: {e}")
    else:
        raise HTTPException(400, "Formato no soportado. Sube un archivo .pdf o .docx")

    if not text:
        raise HTTPException(422, "El documento no contiene texto extraíble")

    return {"filename": filename, "text": text, "chars": len(text), "pages": text.count("\n\n") + 1}


@app.post("/api/generate-doc")
async def generate_document(req: DocRequest):
    """Convert agent output text into a downloadable PDF or DOCX file."""
    safe_name = req.filename.strip() or "reporte"

    if req.format == "pdf":
        data = generate_pdf_bytes(req.content, safe_name)
        return Response(
            content=data,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
        )
    elif req.format == "docx":
        data = generate_docx_bytes(req.content, safe_name)
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )
    else:
        raise HTTPException(400, "Formato no válido. Usa 'pdf' o 'docx'.")


@app.post("/api/run/stream")
async def run_agent_stream(req: RunRequest):
    agent = AGENTS_BY_ID.get(req.agent_id)
    if not agent:
        return {"error": "Agent not found"}

    user_content = req.task
    if req.document_context:
        # Limit doc to first 8 000 chars to stay within token budget
        doc_snippet = req.document_context[:8000]
        user_content = (
            f"Documento adjunto (texto extraído):\n{doc_snippet}\n\n---\n\nTarea:\n{req.task}"
        )
    if req.context:
        user_content = f"Contexto previo:\n{req.context}\n\n{user_content}"

    def generate():
        stream = groq_client.chat.completions.create(
            model=MODEL,
            max_tokens=4000,
            stream=True,
            messages=[
                {"role": "system", "content": agent["system"]},
                {"role": "user", "content": user_content},
            ],
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content
            if delta:
                yield f"data: {json.dumps({'content': delta})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/orchestrate/stream")
async def orchestrate_stream(req: OrchestrationRequest):
    async def generate():
        if req.mode == "sequential":
            context = ""
            for i, step in enumerate(req.steps):
                agent = AGENTS_BY_ID.get(step.agent_id)
                if not agent:
                    continue

                header = json.dumps({
                    "type": "step_start",
                    "step": i + 1,
                    "total": len(req.steps),
                    "agent": agent["name"],
                    "icon": agent["icon"],
                })
                yield f"data: {header}\n\n"

                user_content = step.task
                if context:
                    user_content = f"Contexto del paso anterior:\n{context}\n\nTu tarea:\n{step.task}"

                loop = asyncio.get_event_loop()

                def run_sync():
                    result = []
                    stream = groq_client.chat.completions.create(
                        model=MODEL,
                        max_tokens=2000,
                        stream=True,
                        messages=[
                            {"role": "system", "content": agent["system"]},
                            {"role": "user", "content": user_content},
                        ],
                    )
                    for chunk in stream:
                        delta = chunk.choices[0].delta.content
                        if delta:
                            result.append(delta)
                    return "".join(result)

                step_output = await loop.run_in_executor(None, run_sync)
                context = step_output

                yield f"data: {json.dumps({'type': 'step_result', 'step': i+1, 'content': step_output})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"

        elif req.mode == "parallel":
            results = {}

            async def run_agent(step, idx):
                agent = AGENTS_BY_ID.get(step.agent_id)
                if not agent:
                    return

                loop = asyncio.get_event_loop()

                def run_sync():
                    out = []
                    stream = groq_client.chat.completions.create(
                        model=MODEL,
                        max_tokens=2000,
                        stream=True,
                        messages=[
                            {"role": "system", "content": agent["system"]},
                            {"role": "user", "content": step.task},
                        ],
                    )
                    for chunk in stream:
                        delta = chunk.choices[0].delta.content
                        if delta:
                            out.append(delta)
                    return "".join(out)

                output = await loop.run_in_executor(None, run_sync)
                results[idx] = {"agent": agent["name"], "icon": agent["icon"], "content": output}

            tasks = [run_agent(step, i) for i, step in enumerate(req.steps)]
            await asyncio.gather(*tasks)

            for idx in sorted(results.keys()):
                r = results[idx]
                yield f"data: {json.dumps({'type': 'parallel_result', 'step': idx+1, **r})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


if __name__ == "__main__":
    import uvicorn
    print("\nAgents Hub corriendo en http://localhost:8000\n")
    uvicorn.run("server:app", host="0.0.0.0", port=8000, reload=True)
